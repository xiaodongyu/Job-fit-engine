"""RAG utilities: parsing, chunking, embedding, FAISS index management."""
import os
import re
import json
import uuid
import threading
import numpy as np
import faiss
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
from typing import Optional

from models import (
    UploadStatus,
    EvidenceChunk,
    ExtractionResult,
    ExtractedSkill,
    ExtractedExperience,
    ExperienceItem,
    ClusteredGroup,
)
from prompts import (
    EXTRACT_SYSTEM,
    EXTRACT_SCHEMA,
    build_extract_prompt,
    CLUSTER_SYSTEM,
    CLUSTER_SCHEMA,
    build_cluster_prompt,
    CLUSTER_LABELS,
    RESUME_STRUCTURE_SYSTEM,
    RESUME_STRUCTURE_SCHEMA,
    build_resume_structure_prompt,
    # Two-pass parsing
    RESUME_SEGMENT_SYSTEM,
    RESUME_SEGMENT_SCHEMA,
    build_resume_segment_prompt,
    RESUME_MAP_SYSTEM,
    build_resume_map_prompt,
)

import gemini_client

# === GT role-fit constants (role_percentage_comp_method.md) ===
TIER_WEIGHT = {1: 1.0, 2: 0.6, 3: 0.3}
OWNERSHIP_MULT = {
    "primary": 1.0,
    "parallel": 0.8,
    "earlier_career": 0.7,
    "add_on": 0.6,
    "coursework": 0.4,
}

# === Config ===
DATA_DIR = Path("data")
JD_INDEX_PATH = DATA_DIR / "jd.index"
JD_META_PATH = DATA_DIR / "jd_meta.json"

def get_chunk_size() -> int:
    return int(os.getenv("CHUNK_SIZE", "800"))

def get_chunk_overlap() -> int:
    return int(os.getenv("CHUNK_OVERLAP", "120"))

def get_top_k() -> int:
    return int(os.getenv("TOP_K", "6"))

# === Upload tracking ===
UPLOAD_STATUS: dict[str, dict] = {}  # upload_id -> {status, detail, session_id}

# === Thread pool for background processing ===
executor = ThreadPoolExecutor(max_workers=2)

# Per-session lock: only one processing job per session at a time
SESSION_LOCKS: dict[str, threading.Lock] = {}
_lock_mu = threading.Lock()


def _session_lock(session_id: str) -> threading.Lock:
    with _lock_mu:
        if session_id not in SESSION_LOCKS:
            SESSION_LOCKS[session_id] = threading.Lock()
        return SESSION_LOCKS[session_id]


def ensure_data_dir():
    """Ensure data directory exists."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)


def get_session_dir(session_id: str) -> Path:
    """Get session-specific directory."""
    p = DATA_DIR / "sessions" / session_id
    p.mkdir(parents=True, exist_ok=True)
    return p


# === Parsing ===
def parse_file(file_bytes: bytes, filename: str) -> str:
    """Parse uploaded file to text. Best-effort for PDF/DOCX/TXT."""
    ext = filename.lower().split('.')[-1] if '.' in filename else 'txt'
    
    if ext == 'txt':
        return file_bytes.decode('utf-8', errors='ignore')
    
    elif ext == 'pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            
            # First try: extract text directly
            for page in doc:
                page_text = page.get_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text)
            
            result = "\n".join(text_parts)
            if result.strip():
                doc.close()
                return result
            
            # Second try: OCR for image-based/scanned PDFs
            try:
                import pytesseract
                from PIL import Image
                from io import BytesIO
                import os
                
                # Set Tesseract path for Windows if not in PATH
                if os.name == 'nt':  # Windows
                    tesseract_paths = [
                        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                    ]
                    for tpath in tesseract_paths:
                        if os.path.exists(tpath):
                            pytesseract.pytesseract.tesseract_cmd = tpath
                            break
                
                ocr_text_parts = []
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # Render page to image at 150 DPI for good OCR quality
                    mat = fitz.Matrix(150/72, 150/72)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img = Image.open(BytesIO(img_data))
                    
                    # Run OCR
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text and ocr_text.strip():
                        ocr_text_parts.append(ocr_text)
                
                doc.close()
                result = "\n".join(ocr_text_parts)
                if result.strip():
                    return result
                raise ValueError("PDF appears to be empty (OCR found no text)")
                
            except ImportError:
                doc.close()
                raise ValueError("PDF is image-based but OCR not available. Install Tesseract OCR and run: pip install pytesseract Pillow")
            except Exception as ocr_error:
                doc.close()
                raise ValueError(f"OCR failed on image-based PDF: {ocr_error}")
                
        except ImportError as e:
            raise ValueError(f"PyMuPDF not installed. Run: pip install PyMuPDF. Error: {e}")
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {e}")
    
    elif ext == 'docx':
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(file_bytes))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            result = "\n".join(text_parts)
            if result.strip():
                return result
            raise ValueError("DOCX appears to be empty (no extractable text)")
        except ImportError as e:
            raise ValueError(f"python-docx not installed. Run: pip install python-docx. Error: {e}")
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {e}")
    
    elif ext == 'doc':
        # .doc format requires different library (not supported in MVP)
        raise ValueError(".doc format not supported. Please convert to .docx or .pdf")
    
    else:
        # Try to decode as text
        try:
            result = file_bytes.decode('utf-8', errors='ignore')
            if result.strip():
                return result
            raise ValueError("File appears to be empty")
        except Exception as e:
            raise ValueError(f"Failed to read file as text: {e}")


# === Chunking ===
def chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> list[str]:
    """Split text into overlapping chunks."""
    if chunk_size is None:
        chunk_size = get_chunk_size()
    if overlap is None:
        overlap = get_chunk_overlap()
    
    if not text:
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


# === FAISS Index Management ===
def create_faiss_index(dim: int) -> faiss.IndexFlatIP:
    """Create a FAISS index for inner product (cosine similarity with normalized vectors)."""
    return faiss.IndexFlatIP(dim)


def save_faiss_index(index: faiss.IndexFlatIP, path: Path):
    """Save FAISS index to disk."""
    faiss.write_index(index, str(path))


def load_faiss_index(path: Path) -> Optional[faiss.IndexFlatIP]:
    """Load FAISS index from disk."""
    if path.exists():
        return faiss.read_index(str(path))
    return None


def save_metadata(meta: list[dict], path: Path):
    """Save chunk metadata to JSON."""
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(meta, f, ensure_ascii=False, indent=2)


def load_metadata(path: Path) -> list[dict]:
    """Load chunk metadata from JSON."""
    if path.exists():
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []


RESUME_RAW_FILENAME = "resume_raw.txt"
RESUME_STRUCTURED_FILENAME = "resume_structured.json"
RESUME_BLOCKS_FILENAME = "resume_blocks.json"
RESUME_SEGMENT_PROMPT_FILENAME = "resume_segment_prompt.txt"
RESUME_SEGMENT_RAW_FILENAME = "resume_segment_raw.txt"
RESUME_SEGMENT_PARSED_FILENAME = "resume_segment_parsed.json"
RESUME_MAP_PROMPT_FILENAME = "resume_map_prompt.txt"
RESUME_MAP_RAW_FILENAME = "resume_map_raw.txt"
RESUME_MAP_PARSED_FILENAME = "resume_map_parsed.json"
RESUME_PARSE_TRACE_FILENAME = "resume_parse_trace.json"
EXTRACTION_FILENAME = "extraction.json"
CLUSTERS_FILENAME = "clusters.json"


def _resume_raw_path(session_id: str) -> Path:
    return get_session_dir(session_id) / RESUME_RAW_FILENAME


def _extraction_path(session_id: str) -> Path:
    return get_session_dir(session_id) / EXTRACTION_FILENAME


def _clusters_path(session_id: str) -> Path:
    return get_session_dir(session_id) / CLUSTERS_FILENAME


def _resume_structured_path(session_id: str) -> Path:
    return get_session_dir(session_id) / RESUME_STRUCTURED_FILENAME


def _resume_blocks_path(session_id: str) -> Path:
    return get_session_dir(session_id) / RESUME_BLOCKS_FILENAME


def _resume_debug_path(session_id: str, filename: str) -> Path:
    return get_session_dir(session_id) / filename


def _save_text_checkpoint(session_id: str, filename: str, text: str) -> None:
    try:
        _resume_debug_path(session_id, filename).write_text(text or "", encoding="utf-8")
    except Exception:
        pass


def _save_json_checkpoint(session_id: str, filename: str, data) -> None:
    try:
        with open(_resume_debug_path(session_id, filename), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def save_resume_raw(session_id: str, text: str) -> None:
    path = _resume_raw_path(session_id)
    path.write_text(text, encoding="utf-8")


def load_resume_raw(session_id: str) -> Optional[str]:
    path = _resume_raw_path(session_id)
    if path.exists():
        return path.read_text(encoding="utf-8")
    return None


def save_resume_structured(session_id: str, data: dict) -> None:
    path = _resume_structured_path(session_id)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_resume_structured(session_id: str) -> Optional[dict]:
    path = _resume_structured_path(session_id)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def save_resume_blocks(session_id: str, data: dict) -> None:
    """Save intermediate segmentation blocks for debugging."""
    path = _resume_blocks_path(session_id)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_resume_blocks(session_id: str) -> Optional[dict]:
    """Load intermediate segmentation blocks."""
    path = _resume_blocks_path(session_id)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def _clean_list(value) -> list[str]:
    if not value:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value).strip()] if str(value).strip() else []


def _attach_block_ids(items: list[dict], prefix: str) -> list[dict]:
    out = []
    for item in items:
        block_id = item.get("block_id") or f"{prefix}_{uuid.uuid4().hex[:8]}"
        item["block_id"] = block_id
        out.append(item)
    return out


def _coerce_structured(content: dict) -> dict:
    experiences = content.get("experiences") if isinstance(content.get("experiences"), list) else []
    projects = content.get("projects") if isinstance(content.get("projects"), list) else []
    education = content.get("education") if isinstance(content.get("education"), list) else []

    def normalize_exp(item: dict) -> dict:
        return {
            "block_id": item.get("block_id"),
            "company": item.get("company"),
            "title": item.get("title"),
            "location": item.get("location"),
            "start_date": item.get("start_date"),
            "end_date": item.get("end_date"),
            "bullets": _clean_list(item.get("bullets")),
            "skills_tags": _clean_list(item.get("skills_tags")),
            "ownership": item.get("ownership"),
        }

    def normalize_proj(item: dict) -> dict:
        return {
            "block_id": item.get("block_id"),
            "name": item.get("name"),
            "role": item.get("role"),
            "location": item.get("location"),
            "start_date": item.get("start_date"),
            "end_date": item.get("end_date"),
            "bullets": _clean_list(item.get("bullets")),
            "skills_tags": _clean_list(item.get("skills_tags")),
            "ownership": item.get("ownership"),
        }

    def normalize_edu(item: dict) -> dict:
        return {
            "block_id": item.get("block_id"),
            "school": item.get("school"),
            "degree": item.get("degree"),
            "field": item.get("field"),
            "location": item.get("location"),
            "start_date": item.get("start_date"),
            "end_date": item.get("end_date"),
            "gpa": item.get("gpa"),
            "bullets": _clean_list(item.get("bullets")),
        }

    structured = {
        "experiences": _attach_block_ids([normalize_exp(i) for i in experiences], "exp"),
        "projects": _attach_block_ids([normalize_proj(i) for i in projects], "proj"),
        "education": _attach_block_ids([normalize_edu(i) for i in education], "edu"),
    }
    return structured


def _strip_noise_sections(lines: list[str]) -> list[str]:
    """Remove known non-resume sections like role-fit scores."""
    cleaned = []
    skip = False
    for ln in lines:
        upper = ln.upper()
        if upper.startswith("GROUND-TRUTH ROLE FIT") or upper.startswith("ROLE FIT"):
            skip = True
            continue
        if skip and re.match(r"^-?\s*(MLE|SWE|DS|QD|QR)\s*:", ln, re.IGNORECASE):
            continue
        if skip and ln.strip() == "":
            skip = False
            continue
        if skip and upper.startswith("SUMMARY"):
            skip = False
        if skip:
            continue
        cleaned.append(ln)
    return cleaned


def _looks_like_experience_header(line: str) -> bool:
    # Examples: "Senior ML Engineer — Company | 2020–Present"
    return bool(re.search(r"—.*\|\s*\d{4}", line)) or bool(re.search(r"\b\d{4}\s*[–-]\s*(Present|\d{4})", line, re.IGNORECASE))


def _heuristic_structured(text: str) -> dict:
    """Parse resume text into structured blocks using heuristics.
    
    Detects section headers and groups content appropriately.
    """
    import re
    raw_lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    lines = _strip_noise_sections(raw_lines)
    
    # Section header patterns
    SECTION_PATTERNS = {
        "education": re.compile(r"^(EDUCATION|ACADEMIC|DEGREE)", re.IGNORECASE),
        "experience": re.compile(r"^(EXPERIENCE|EMPLOYMENT|WORK|CAREER|PROFESSIONAL)", re.IGNORECASE),
        "projects": re.compile(r"^(PROJECT|PERSONAL PROJECT|ACADEMIC PROJECT)", re.IGNORECASE),
        "skills": re.compile(r"^(SKILL|TECHNICAL SKILL|COMPETENC)", re.IGNORECASE),
    }
    
    experiences = []
    projects = []
    education = []
    current_section = None  # "education", "experience", "projects", or None
    current_block = None
    
    def _save_current_block():
        nonlocal current_block
        if current_block is None:
            return
        if current_section == "education":
            education.append(current_block)
        elif current_section == "experience":
            experiences.append(current_block)
        elif current_section == "projects":
            projects.append(current_block)
        current_block = None
    
    def _is_section_header(line: str) -> Optional[str]:
        """Check if line is a section header, return section type or None."""
        for section, pattern in SECTION_PATTERNS.items():
            if pattern.match(line):
                return section
        return None
    
    def _parse_date_from_line(line: str) -> tuple[Optional[str], Optional[str], str]:
        """Extract dates from a line, return (start, end, remaining_text)."""
        date_pattern = re.compile(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|"
            r"\d{4}[-/]\d{2}|\d{4})"
            r"\s*[-–—to]+\s*"
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|"
            r"\d{4}[-/]\d{2}|\d{4}|Present|Current)",
            re.IGNORECASE
        )
        match = date_pattern.search(line)
        if match:
            start = match.group(1)
            end = match.group(2)
            remaining = date_pattern.sub("", line).strip(" ,|-–")
            return start, end, remaining
        return None, None, line
    
    for ln in lines:
        # Check for section headers first
        section_type = _is_section_header(ln)
        if section_type:
            _save_current_block()
            current_section = section_type
            continue
        
        # Skip if we haven't identified a section yet
        if current_section is None:
            # Try to detect based on content
            if _looks_like_experience_header(ln):
                current_section = "experience"
            else:
                continue
        
        # Handle bullet points - add to current block
        if ln.startswith(("-", "•", "*", "◦", "▪")):
            bullet_text = ln.lstrip("-•*◦▪ ").strip()
            if current_block and bullet_text:
                current_block["bullets"].append(bullet_text)
            continue
        
        # Non-bullet line - could be a new block header or continuation
        start_date, end_date, remaining = _parse_date_from_line(ln)
        
        if current_section == "education":
            # Education: look for school/degree patterns
            # New education block if we see a school name (contains "University", "College", etc.)
            # or if dates are present
            is_new_block = (
                start_date is not None or
                re.search(r"(University|College|Institute|School|Academy)", ln, re.IGNORECASE) or
                re.search(r"(Bachelor|Master|Ph\.?D|B\.?S\.?|M\.?S\.?|M\.?A\.?|B\.?A\.?)", ln, re.IGNORECASE)
            )
            if is_new_block:
                _save_current_block()
                current_block = {
                    "block_id": f"edu_{uuid.uuid4().hex[:8]}",
                    "school": None,
                    "degree": None,
                    "field": None,
                    "location": None,
                    "start_date": start_date,
                    "end_date": end_date,
                    "gpa": None,
                    "bullets": [],
                }
                # Try to extract school name
                school_match = re.search(r"([\w\s]+(?:University|College|Institute|School|Academy)[\w\s]*)", ln, re.IGNORECASE)
                if school_match:
                    current_block["school"] = school_match.group(1).strip()
                # Try to extract degree
                degree_match = re.search(r"(Bachelor|Master|Ph\.?D|B\.?S\.?|M\.?S\.?|M\.?A\.?|B\.?A\.?)[\s\w]*", ln, re.IGNORECASE)
                if degree_match:
                    current_block["degree"] = degree_match.group(0).strip()
            elif current_block:
                # Continuation - add as bullet or update fields
                if ln.lower().startswith("gpa"):
                    gpa_match = re.search(r"(\d+\.?\d*)", ln)
                    if gpa_match:
                        current_block["gpa"] = gpa_match.group(1)
                    else:
                        current_block["bullets"].append(ln)
                else:
                    current_block["bullets"].append(ln)
                    
        elif current_section == "experience":
            # Experience: look for company/title patterns
            is_new_block = _looks_like_experience_header(ln) or start_date is not None
            if is_new_block:
                _save_current_block()
                current_block = {
                    "block_id": f"exp_{uuid.uuid4().hex[:8]}",
                    "company": None,
                    "title": remaining if remaining else ln,
                    "location": None,
                    "start_date": start_date,
                    "end_date": end_date,
                    "bullets": [],
                    "skills_tags": [],
                    "ownership": None,
                }
            elif current_block:
                # Continuation - add as bullet
                current_block["bullets"].append(ln)
                
        elif current_section == "projects":
            # Projects: similar to experience
            is_new_block = start_date is not None or (not ln.startswith(" ") and len(ln) < 100)
            if is_new_block and (start_date or not current_block):
                _save_current_block()
                current_block = {
                    "block_id": f"proj_{uuid.uuid4().hex[:8]}",
                    "name": remaining if remaining else ln,
                    "role": None,
                    "location": None,
                    "start_date": start_date,
                    "end_date": end_date,
                    "bullets": [],
                    "skills_tags": [],
                    "ownership": None,
                }
            elif current_block:
                current_block["bullets"].append(ln)
    
    # Save final block
    _save_current_block()
    
    # Fallback: if nothing was parsed, treat all bullets as one experience
    if not experiences and not education and not projects:
        bullets = [ln.lstrip("-•* ").strip() for ln in lines if ln.startswith(("-", "•", "*"))]
        if not bullets:
            bullets = lines[:10]
        experiences = [{
            "block_id": f"exp_{uuid.uuid4().hex[:8]}",
            "company": None,
            "title": None,
            "location": None,
            "start_date": None,
            "end_date": None,
            "bullets": bullets,
            "skills_tags": [],
            "ownership": None,
        }]
    
    return {
        "experiences": experiences,
        "projects": projects,
        "education": education,
    }


# === Two-Pass Resume Parsing ===

def segment_resume_blocks(text: str) -> dict:
    """Pass 1: Segment preprocessed resume text into blocks.
    
    Returns: {"blocks": [{"section": ..., "header_lines": [...], ...}]}
    """
    prompt = build_resume_segment_prompt(text)
    result = gemini_client.generate(RESUME_SEGMENT_SYSTEM, prompt, RESUME_SEGMENT_SCHEMA)
    content = result.get("content") or {}
    if isinstance(content, str):
        content = {}
    
    blocks = content.get("blocks")
    if not blocks or not isinstance(blocks, list):
        return {"blocks": []}
    
    # Validate each block has required fields
    validated = []
    for b in blocks:
        if not isinstance(b, dict):
            continue
        validated.append({
            "section": b.get("section", "other"),
            "header_lines": b.get("header_lines") or [],
            "meta_lines": b.get("meta_lines") or [],
            "bullet_lines": b.get("bullet_lines") or [],
            "raw_lines": b.get("raw_lines") or [],
        })
    
    return {"blocks": validated}


def map_blocks_to_structured(blocks_data: dict) -> dict:
    """Pass 2: Map segmented blocks to final RESUME_STRUCTURE_SCHEMA.
    
    Takes output from segment_resume_blocks() and returns structured resume.
    """
    blocks = blocks_data.get("blocks", [])
    if not blocks:
        return {"experiences": [], "projects": [], "education": []}
    
    prompt = build_resume_map_prompt(blocks_data)
    result = gemini_client.generate(RESUME_MAP_SYSTEM, prompt, RESUME_STRUCTURE_SCHEMA)
    content = result.get("content") or {}
    if isinstance(content, str):
        content = {}
    
    return _coerce_structured(content)


def extract_resume_structure(text: str, session_id: str = None) -> dict:
    """Extract structured experiences/projects/education from resume text.
    
    Uses two-pass parsing:
      1. Segment text into blocks (preserving raw lines)
      2. Map blocks to final schema
    
    Falls back to single-pass or heuristic parsing if two-pass fails.
    
    Args:
        text: Raw resume text
        session_id: Optional session ID for saving intermediate blocks
    
    Returns:
        Structured resume dict with experiences, projects, education arrays.
    """
    cleaned_text = _preprocess_resume_text(text)

    trace: dict = {
        "session_id": session_id,
        "path_used": None,
        "two_pass": {
            "segment": {"ok": False, "error": None, "blocks_count": 0},
            "map": {"ok": False, "error": None, "experiences": 0, "projects": 0, "education": 0},
        },
        "single_pass": {"ok": False, "error": None, "experiences": 0, "projects": 0, "education": 0},
        "heuristic": {"used": False},
    }

    # Try two-pass approach
    try:
        # Pass 1: Segmentation
        seg_prompt = build_resume_segment_prompt(cleaned_text)
        if session_id:
            _save_text_checkpoint(session_id, RESUME_SEGMENT_PROMPT_FILENAME, seg_prompt)
        seg_result = gemini_client.generate(RESUME_SEGMENT_SYSTEM, seg_prompt, RESUME_SEGMENT_SCHEMA)
        if session_id:
            _save_text_checkpoint(session_id, RESUME_SEGMENT_RAW_FILENAME, seg_result.get("raw") or "")
            _save_json_checkpoint(session_id, RESUME_SEGMENT_PARSED_FILENAME, seg_result.get("content"))

        seg_content = seg_result.get("content") or {}
        if isinstance(seg_content, str):
            seg_content = {}
        blocks = seg_content.get("blocks")
        if not blocks or not isinstance(blocks, list):
            blocks_data = {"blocks": []}
        else:
            # Validate shape (same logic as segment_resume_blocks)
            validated = []
            for b in blocks:
                if not isinstance(b, dict):
                    continue
                validated.append({
                    "section": b.get("section", "other"),
                    "header_lines": b.get("header_lines") or [],
                    "meta_lines": b.get("meta_lines") or [],
                    "bullet_lines": b.get("bullet_lines") or [],
                    "raw_lines": b.get("raw_lines") or [],
                })
            blocks_data = {"blocks": validated}

        trace["two_pass"]["segment"]["ok"] = True
        trace["two_pass"]["segment"]["blocks_count"] = len(blocks_data.get("blocks") or [])

        # Save blocks for debugging
        if session_id:
            save_resume_blocks(session_id, blocks_data)

        # Pass 2 only if we have blocks
        if blocks_data.get("blocks"):
            map_prompt = build_resume_map_prompt(blocks_data)
            if session_id:
                _save_text_checkpoint(session_id, RESUME_MAP_PROMPT_FILENAME, map_prompt)
            map_result = gemini_client.generate(RESUME_MAP_SYSTEM, map_prompt, RESUME_STRUCTURE_SCHEMA)
            if session_id:
                _save_text_checkpoint(session_id, RESUME_MAP_RAW_FILENAME, map_result.get("raw") or "")
                _save_json_checkpoint(session_id, RESUME_MAP_PARSED_FILENAME, map_result.get("content"))

            map_content = map_result.get("content") or {}
            if isinstance(map_content, str):
                map_content = {}
            structured = _coerce_structured(map_content)
            trace["two_pass"]["map"]["experiences"] = len(structured.get("experiences") or [])
            trace["two_pass"]["map"]["projects"] = len(structured.get("projects") or [])
            trace["two_pass"]["map"]["education"] = len(structured.get("education") or [])
            trace["two_pass"]["map"]["ok"] = True

            has_blocks = any(structured.get(k) for k in ("experiences", "projects", "education"))
            if has_blocks:
                trace["path_used"] = "two_pass"
                if session_id:
                    _save_json_checkpoint(session_id, RESUME_PARSE_TRACE_FILENAME, trace)
                return structured
    except Exception as e:
        trace["two_pass"]["segment"]["error"] = str(e)
        trace["two_pass"]["map"]["error"] = str(e)
        # Fall through to single-pass
    
    # Fallback: Single-pass approach (original method)
    try:
        prompt = build_resume_structure_prompt(cleaned_text)
        result = gemini_client.generate(RESUME_STRUCTURE_SYSTEM, prompt, RESUME_STRUCTURE_SCHEMA)
        content = result.get("content") or {}
        if isinstance(content, str):
            content = {}
        structured = _coerce_structured(content)
        has_blocks = any(structured.get(k) for k in ("experiences", "projects", "education"))
        if has_blocks:
            trace["single_pass"]["ok"] = True
            trace["single_pass"]["experiences"] = len(structured.get("experiences") or [])
            trace["single_pass"]["projects"] = len(structured.get("projects") or [])
            trace["single_pass"]["education"] = len(structured.get("education") or [])
            trace["path_used"] = "single_pass"
            if session_id:
                _save_json_checkpoint(session_id, RESUME_PARSE_TRACE_FILENAME, trace)
            return structured
    except Exception:
        trace["single_pass"]["error"] = "single_pass_exception"
        # Fall through to heuristic
    
    # Final fallback: Heuristic parsing
    trace["heuristic"]["used"] = True
    trace["path_used"] = "heuristic"
    if session_id:
        _save_json_checkpoint(session_id, RESUME_PARSE_TRACE_FILENAME, trace)
    return _heuristic_structured(cleaned_text)


def _preprocess_resume_text(text: str) -> str:
    """Normalize tab-separated resume layouts for better parsing.

    Converts patterns like:
      Company\\tLocation
      Title\\tDates
    Into:
      Company | Location
      Title | Dates
    """
    import re

    lines = text.splitlines()
    processed: list[str] = []

    for line in lines:
        # Replace tabs with pipe separator (more explicit for LLM)
        if "\t" in line:
            # Normalize multiple tabs/spaces to single pipe
            line = re.sub(r"\t+\s*", " | ", line)
        # Clean up excessive whitespace
        line = re.sub(r"  +", " ", line).strip()
        processed.append(line)

    return "\n".join(processed)


def _format_date_range(start: Optional[str], end: Optional[str]) -> Optional[str]:
    if start and end:
        return f"{start}–{end}"
    return start or end


def _build_header(kind: str, block: dict) -> str:
    if kind == "experience":
        parts = [block.get("company"), block.get("title"), block.get("location")]
        date_part = _format_date_range(block.get("start_date"), block.get("end_date"))
        if date_part:
            parts.append(date_part)
        return " | ".join([p for p in parts if p]) or "EXPERIENCE"
    if kind == "project":
        parts = [block.get("name"), block.get("role"), block.get("location")]
        date_part = _format_date_range(block.get("start_date"), block.get("end_date"))
        if date_part:
            parts.append(date_part)
        return " | ".join([p for p in parts if p]) or "PROJECT"
    if kind == "education":
        parts = [block.get("school"), block.get("degree"), block.get("field"), block.get("location")]
        date_part = _format_date_range(block.get("start_date"), block.get("end_date"))
        if date_part:
            parts.append(date_part)
        return " | ".join([p for p in parts if p]) or "EDUCATION"
    return "RESUME"


def _chunk_block_with_header(
    block: dict,
    kind: str,
    max_chars: int,
    overlap: int,
) -> tuple[list[str], list[dict]]:
    header = _build_header(kind, block)
    bullets = _clean_list(block.get("bullets"))
    tags = _clean_list(block.get("skills_tags"))
    ownership = block.get("ownership")

    body_lines = []
    if ownership:
        body_lines.append(f"Ownership: {ownership}")
    if tags:
        body_lines.append("Tags: " + ", ".join(tags))
    for b in bullets:
        body_lines.append(f"- {b}")

    body = "\n".join(body_lines).strip()
    block_id = block.get("block_id") or f"{kind}_{uuid.uuid4().hex[:8]}"

    chunks = []
    meta = []
    if not body:
        chunk_id = f"{block_id}__{kind}__0"
        chunks.append(header)
        meta.append({
            "chunk_id": chunk_id,
            "text": header,
            "source_type": "resume",
            "session_id": None,
            "block_id": block_id,
            "block_type": kind,
            "sub_index": 0,
            "header": header,
        })
        return chunks, meta

    available = max(200, max_chars - len(header) - 2)
    start = 0
    idx = 0
    while start < len(body):
        end = min(len(body), start + available)
        segment = body[start:end].strip()
        text = f"{header}\n{segment}"
        chunk_id = f"{block_id}__{kind}__{idx}"
        chunks.append(text)
        meta.append({
            "chunk_id": chunk_id,
            "text": text,
            "source_type": "resume",
            "session_id": None,
            "block_id": block_id,
            "block_type": kind,
            "sub_index": idx,
            "header": header,
        })
        idx += 1
        if end >= len(body):
            break
        start = max(0, end - overlap)

    return chunks, meta


def chunk_structured_resume(structured: dict) -> tuple[list[str], list[dict]]:
    max_chars = int(os.getenv("BLOCK_CHUNK_SIZE", "1200"))
    overlap = int(os.getenv("BLOCK_CHUNK_OVERLAP", "150"))
    chunks: list[str] = []
    meta: list[dict] = []

    for exp in structured.get("experiences", []) or []:
        c, m = _chunk_block_with_header(exp, "experience", max_chars, overlap)
        chunks.extend(c)
        meta.extend(m)
    for proj in structured.get("projects", []) or []:
        c, m = _chunk_block_with_header(proj, "project", max_chars, overlap)
        chunks.extend(c)
        meta.extend(m)
    for edu in structured.get("education", []) or []:
        c, m = _chunk_block_with_header(edu, "education", max_chars, overlap)
        chunks.extend(c)
        meta.extend(m)

    return chunks, meta


def save_extraction(session_id: str, data: dict) -> None:
    path = _extraction_path(session_id)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_extraction(session_id: str) -> Optional[dict]:
    path = _extraction_path(session_id)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def save_clusters(session_id: str, data: dict) -> None:
    path = _clusters_path(session_id)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_clusters(session_id: str) -> Optional[dict]:
    path = _clusters_path(session_id)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return None


def run_extraction(meta: list[dict]) -> ExtractionResult:
    """Extract skills and experiences from resume chunks via Gemini."""
    chunks = [{"chunk_id": m["chunk_id"], "text": m["text"]} for m in meta]
    if not chunks:
        return ExtractionResult(skills=[], experiences=[])
    prompt = build_extract_prompt(chunks)
    result = gemini_client.generate(EXTRACT_SYSTEM, prompt, EXTRACT_SCHEMA)
    content = result.get("content") or {}
    if isinstance(content, str):
        content = {}
    skills = [
        ExtractedSkill(name=s.get("name", ""), chunk_ids=s.get("chunk_ids", []))
        for s in content.get("skills", [])
    ]
    experiences = [
        ExtractedExperience(text=e.get("text", ""), chunk_ids=e.get("chunk_ids", []))
        for e in content.get("experiences", [])
    ]
    return ExtractionResult(skills=skills, experiences=experiences)


def run_clustering(extraction: ExtractionResult, meta: list[dict]) -> tuple[list[ClusteredGroup], dict[str, float]]:
    """Classify skills/experiences into MLE/DS/SWE/QR/QD and build ClusteredGroups with evidence."""
    chunk_map = {m["chunk_id"]: m["text"] for m in meta}
    ext_dict = {
        "skills": [{"name": s.name, "chunk_ids": s.chunk_ids} for s in extraction.skills],
        "experiences": [{"text": e.text, "chunk_ids": e.chunk_ids} for e in extraction.experiences],
    }
    prompt = build_cluster_prompt(ext_dict, chunk_map)
    result = gemini_client.generate(CLUSTER_SYSTEM, prompt, CLUSTER_SCHEMA)
    content = result.get("content") or {}
    if isinstance(content, str):
        content = {}
    assignments = content.get("assignments", [])
    valid_cids = ["MLE", "DS", "SWE", "QR", "QD"]

    # GT role_fit_distribution when role_tiers present; else equal-split fallback
    use_gt = any(
        isinstance(a.get("role_tiers"), list) and len(a.get("role_tiers", [])) > 0
        for a in assignments
    )
    weighted: dict[str, float] = {c: 0.0 for c in valid_cids}
    if use_gt:
        for a in assignments:
            role_tiers = a.get("role_tiers") or []
            ownership_mult = OWNERSHIP_MULT.get(a.get("ownership"), 1.0)
            for rt in role_tiers:
                role_id = rt.get("role")
                tier_raw = rt.get("tier")
                tier = int(tier_raw) if isinstance(tier_raw, str) and tier_raw.isdigit() else tier_raw
                if role_id not in valid_cids:
                    continue
                tw = TIER_WEIGHT.get(tier)
                if tw is None:
                    continue
                weighted[role_id] += tw * ownership_mult
    else:
        for a in assignments:
            clusters_list = [x for x in a.get("clusters", []) if x in valid_cids]
            if not clusters_list:
                continue
            w = 1.0 / len(clusters_list)
            for cid in clusters_list:
                weighted[cid] += w
    total_w = sum(weighted.values())
    role_fit_distribution = {c: (weighted[c] / total_w if total_w > 0 else 0.0) for c in valid_cids}

    # Build cluster_id -> {items, chunk_ids for evidence}
    # Derive clusters from role_tiers when present, else use clusters
    clusters_data: dict[str, dict] = {}
    for cid in valid_cids:
        clusters_data[cid] = {"items": [], "chunk_ids": set()}

    for a in assignments:
        item_type = a.get("item_type", "skill")
        item_value = a.get("item_value", "")
        if use_gt and isinstance(a.get("role_tiers"), list) and a["role_tiers"]:
            clusters_list = [rt["role"] for rt in a["role_tiers"] if rt.get("role") in valid_cids]
        else:
            clusters_list = [x for x in a.get("clusters", []) if x in valid_cids]
        chunk_ids = a.get("chunk_ids", [])
        uid = uuid.uuid4().hex[:8]
        item = ExperienceItem(
            id=f"ext_{uid}",
            label=item_type,
            text=item_value,
            source="extraction",
        )
        for cid in clusters_list:
            clusters_data[cid]["items"].append(item)
            for c in chunk_ids:
                clusters_data[cid]["chunk_ids"].add(c)

    out = []
    for cid, label in CLUSTER_LABELS.items():
        data = clusters_data[cid]
        if not data["items"]:
            continue
        evidence = [
            EvidenceChunk(
                chunk_id=cid_,
                text=chunk_map.get(cid_, ""),
                source="resume",
                score=1.0,
            )
            for cid_ in data["chunk_ids"]
            if cid_ in chunk_map
        ]
        out.append(
            ClusteredGroup(
                cluster_id=cid,
                cluster_label=label,
                items=data["items"],
                summary="",
                evidence=evidence,
            )
        )
    return out, role_fit_distribution


# === JD Index (Global) ===
def ingest_jds(items: list[dict]) -> int:
    """
    Ingest JD items into global FAISS index.
    items: list of {title, role, level, text}
    Returns count of chunks added.
    """
    ensure_data_dir()
    
    # Load existing index and metadata
    existing_index = load_faiss_index(JD_INDEX_PATH)
    existing_meta = load_metadata(JD_META_PATH)
    
    all_chunks = []
    all_meta = []
    
    for item in items:
        text = item["text"]
        chunks = chunk_text(text)
        
        for i, chunk in enumerate(chunks):
            chunk_id = f"jd_{item['role']}_{item['level']}_{uuid.uuid4().hex[:8]}_{i}"
            all_chunks.append(chunk)
            all_meta.append({
                "chunk_id": chunk_id,
                "text": chunk,
                "source_type": "jd",
                "doc_id": item.get("title", "unknown"),
                "role": item["role"],
                "level": item["level"],
                "title": item.get("title", "")
            })
    
    if not all_chunks:
        return 0
    
    # Embed chunks
    embeddings = gemini_client.embed_texts(all_chunks)
    embeddings_np = np.array(embeddings, dtype=np.float32)
    
    dim = embeddings_np.shape[1]
    
    # Create or extend index
    if existing_index is None:
        index = create_faiss_index(dim)
    else:
        index = existing_index
    
    index.add(embeddings_np)
    
    # Merge metadata
    combined_meta = existing_meta + all_meta
    
    # Save
    save_faiss_index(index, JD_INDEX_PATH)
    save_metadata(combined_meta, JD_META_PATH)
    
    return len(all_chunks)


def search_jd_index(query: str, role: str = None, top_k: int = None) -> list[EvidenceChunk]:
    """Search global JD index."""
    if top_k is None:
        top_k = get_top_k()
    
    index = load_faiss_index(JD_INDEX_PATH)
    meta = load_metadata(JD_META_PATH)
    
    if index is None or not meta:
        print(f"[DEBUG] search_jd_index: index={index}, meta_len={len(meta) if meta else 0}")
        return []
    
    # Debug: print unique roles in metadata
    unique_roles = set(m.get("role") for m in meta)
    print(f"[DEBUG] search_jd_index: filter_role={role!r} (type={type(role).__name__}), available_roles={unique_roles}, total_chunks={len(meta)}")
    
    # Embed query
    query_vec = gemini_client.embed_single(query)
    query_np = np.array([query_vec], dtype=np.float32)
    
    # Search
    k = min(top_k * 3, index.ntotal)  # Over-fetch for filtering
    scores, indices = index.search(query_np, k)
    
    results = []
    skipped_roles = []
    # Convert enum to string value for comparison
    role_str = role.value if hasattr(role, 'value') else (str(role) if role else None)
    print(f"[DEBUG] role_str after conversion: {role_str!r}")
    
    for score, idx in zip(scores[0], indices[0]):
        if idx < 0 or idx >= len(meta):
            continue
        m = meta[idx]
        meta_role = m.get("role")
        if role_str and meta_role != role_str:
            skipped_roles.append(meta_role)
            continue
        results.append(EvidenceChunk(
            chunk_id=m["chunk_id"],
            text=m["text"],
            source="jd",
            score=float(score)
        ))
        if len(results) >= top_k:
            break
    
    print(f"[DEBUG] search_jd_index: returned {len(results)} chunks, skipped roles: {set(skipped_roles)}")
    return results


# === Resume Index (Per-session) ===
def update_upload_status(upload_id: str, status: UploadStatus, detail: str = ""):
    """Update upload status."""
    if upload_id in UPLOAD_STATUS:
        UPLOAD_STATUS[upload_id]["status"] = status
        UPLOAD_STATUS[upload_id]["detail"] = detail


def get_upload_status(upload_id: str) -> dict:
    """Get upload status."""
    return UPLOAD_STATUS.get(upload_id, {"status": UploadStatus.ERROR, "detail": "Unknown upload_id"})


def _run_full_pipeline(upload_id: str, session_id: str, text: str) -> None:
    """Chunk -> embed -> index -> extract -> cluster. Used by resume upload and add-materials."""
    session_dir = get_session_dir(session_id)
    index_path = session_dir / "resume.index"
    meta_path = session_dir / "resume_meta.json"

    update_upload_status(upload_id, UploadStatus.PARSING, "Extracting structured resume blocks (two-pass)...")
    structured = extract_resume_structure(text, session_id=session_id)
    save_resume_structured(session_id, structured)

    update_upload_status(upload_id, UploadStatus.CHUNKING, "Chunking structured blocks...")
    chunks, meta = chunk_structured_resume(structured)
    if not chunks:
        update_upload_status(upload_id, UploadStatus.CHUNKING, "Falling back to sliding-window chunking...")
        chunks = chunk_text(text)
        meta = [
            {
                "chunk_id": f"resume_{session_id}_{i}",
                "text": chunk,
                "source_type": "resume",
                "session_id": session_id,
                "chunk_index": i,
            }
            for i, chunk in enumerate(chunks)
        ]
    else:
        for i, m in enumerate(meta):
            m["session_id"] = session_id
            m["chunk_index"] = i

    update_upload_status(upload_id, UploadStatus.EMBEDDING, f"Embedding {len(chunks)} chunks...")
    embeddings = gemini_client.embed_texts(chunks)
    embeddings_np = np.array(embeddings, dtype=np.float32)

    for i, chunk in enumerate(chunks):
        if i >= len(meta):
            meta.append({
                "chunk_id": f"resume_{session_id}_{i}",
                "text": chunk,
                "source_type": "resume",
                "session_id": session_id,
                "chunk_index": i,
            })

    update_upload_status(upload_id, UploadStatus.INDEXING, "Building search index...")
    dim = embeddings_np.shape[1]
    index = create_faiss_index(dim)
    index.add(embeddings_np)
    save_faiss_index(index, index_path)
    save_metadata(meta, meta_path)

    # Extraction (1b)
    update_upload_status(upload_id, UploadStatus.INDEXING, "Extracting skills and experiences...")
    extraction = run_extraction(meta)
    save_extraction(session_id, extraction.model_dump())

    # Clustering (2)
    update_upload_status(upload_id, UploadStatus.INDEXING, "Clustering into MLE/DS/SWE/QR/QD...")
    clusters, role_fit_distribution = run_clustering(extraction, meta)
    clusters_data = {
        "clusters": [c.model_dump() for c in clusters],
        "total_items": sum(len(c.items) for c in clusters),
        "role_fit_distribution": role_fit_distribution,
    }
    save_clusters(session_id, clusters_data)

    update_upload_status(upload_id, UploadStatus.READY, f"Indexed {len(chunks)} chunks; {len(clusters)} clusters")
    return


def process_resume_background(upload_id: str, session_id: str, text: str) -> None:
    """Background task: save raw -> run full pipeline (chunk, embed, index, extract, cluster)."""
    lock = _session_lock(session_id)
    with lock:
        try:
            save_resume_raw(session_id, text)
            _run_full_pipeline(upload_id, session_id, text)
        except Exception as e:
            update_upload_status(upload_id, UploadStatus.ERROR, str(e))


def start_resume_processing(session_id: str, text: str) -> str:
    """
    Start background processing of resume.
    Returns upload_id.
    """
    upload_id = uuid.uuid4().hex[:12]
    UPLOAD_STATUS[upload_id] = {
        "status": UploadStatus.PARSING,
        "detail": "Processing uploaded content...",
        "session_id": session_id,
    }
    executor.submit(process_resume_background, upload_id, session_id, text)
    return upload_id


def _add_materials_worker(upload_id: str, session_id: str, new_text: str) -> None:
    """Merge new_text into session resume_raw and re-run full pipeline."""
    lock = _session_lock(session_id)
    with lock:
        try:
            raw = load_resume_raw(session_id)
            if not raw or not raw.strip():
                update_upload_status(upload_id, UploadStatus.ERROR, "No existing resume to merge. Upload resume first.")
                return
            merged = raw.strip() + "\n\n" + new_text.strip()
            save_resume_raw(session_id, merged)
            _run_full_pipeline(upload_id, session_id, merged)
        except Exception as e:
            update_upload_status(upload_id, UploadStatus.ERROR, str(e))


def start_add_materials(session_id: str, new_text: str) -> str:
    """
    Merge new_text into session resume and re-run pipeline (chunk, extract, cluster, index).
    Returns upload_id for status polling. Session must have processed resume (resume_raw.txt) already.
    """
    if not load_resume_raw(session_id):
        raise ValueError(f"Session {session_id} has no resume. Upload resume first.")
    upload_id = uuid.uuid4().hex[:12]
    UPLOAD_STATUS[upload_id] = {
        "status": UploadStatus.PARSING,
        "detail": "Merging materials and re-processing...",
        "session_id": session_id,
    }
    executor.submit(_add_materials_worker, upload_id, session_id, new_text)
    return upload_id


def search_resume_index(session_id: str, query: str, top_k: int = None) -> list[EvidenceChunk]:
    """Search session's resume index."""
    if top_k is None:
        top_k = get_top_k()
    
    session_dir = get_session_dir(session_id)
    index_path = session_dir / "resume.index"
    meta_path = session_dir / "resume_meta.json"
    
    index = load_faiss_index(index_path)
    meta = load_metadata(meta_path)
    
    if index is None or not meta:
        return []
    
    # Embed query
    query_vec = gemini_client.embed_single(query)
    query_np = np.array([query_vec], dtype=np.float32)
    
    # Search
    k = min(top_k, index.ntotal)
    scores, indices = index.search(query_np, k)
    
    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx < 0 or idx >= len(meta):
            continue
        m = meta[idx]
        results.append(EvidenceChunk(
            chunk_id=m["chunk_id"],
            text=m["text"],
            source="resume",
            score=float(score)
        ))
    
    return results


def session_is_ready(session_id: str) -> bool:
    """Check if session's resume index is ready."""
    session_dir = get_session_dir(session_id)
    index_path = session_dir / "resume.index"
    return index_path.exists()


def get_all_resume_chunks(session_id: str) -> list[dict]:
    """
    Retrieve all resume chunks for a session (without search/query).
    Returns list of {chunk_id, text, source_type, session_id, chunk_index}.
    """
    session_dir = get_session_dir(session_id)
    meta_path = session_dir / "resume_meta.json"
    
    if not meta_path.exists():
        return []
    
    return load_metadata(meta_path)


# === Temporary JD processing (for user-pasted JD) ===
def search_temp_jd(jd_text: str, query: str, top_k: int = None) -> list[EvidenceChunk]:
    """
    Create temporary in-memory index for user-pasted JD and search.
    """
    if top_k is None:
        top_k = get_top_k()
    
    # Chunk the JD text
    chunks = chunk_text(jd_text)
    if not chunks:
        return []
    
    # Embed chunks
    embeddings = gemini_client.embed_texts(chunks)
    embeddings_np = np.array(embeddings, dtype=np.float32)
    
    # Build temp index
    dim = embeddings_np.shape[1]
    index = create_faiss_index(dim)
    index.add(embeddings_np)
    
    # Embed query and search
    query_vec = gemini_client.embed_single(query)
    query_np = np.array([query_vec], dtype=np.float32)
    
    k = min(top_k, index.ntotal)
    scores, indices = index.search(query_np, k)
    
    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx < 0 or idx >= len(chunks):
            continue
        results.append(EvidenceChunk(
            chunk_id=f"temp_jd_{idx}",
            text=chunks[idx],
            source="jd",
            score=float(score)
        ))
    
    return results
