"""FastAPI application with all endpoints."""
import uuid
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

load_dotenv()

from models import (
    JDIngestRequest, JDIngestResponse,
    ResumeUploadTextRequest, ResumeUploadResponse, ResumeStatusResponse,
    AddMaterialsRequest,
    AnalyzeFitRequest, AnalyzeFitResponse,
    RoleRecommendation, Requirements, GapAnalysis, Evidence, EvidenceChunk,
    ResumeGenerateRequest, ResumeGenerateResponse, ResumeStructured,
    HealthResponse, UploadStatus,
    ClusterRequest, ClusterResponse, ClusteredGroup, ExperienceItem,
    MatchByClusterRequest, MatchByClusterResponse, ClusterMatch,
    ResumeStructuredResponse, ResumeInputStructured,
)
from prompts import (
    ROLE_FIT_SYSTEM, ROLE_FIT_SCHEMA,
    RESUME_GENERATE_SYSTEM, RESUME_GENERATE_SCHEMA,
    build_fit_prompt, build_generate_prompt,
    MATCH_BY_CLUSTER_SYSTEM, MATCH_BY_CLUSTER_SCHEMA, build_match_by_cluster_prompt,
)
import rag
import gemini_client

app = FastAPI(title="Tech Career Fit Engine", version="0.2.0")

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health", response_model=HealthResponse)
async def health():
    """Health check endpoint."""
    return HealthResponse(status="ok", message="Tech Career Fit Engine is running")


@app.post("/jd/ingest", response_model=JDIngestResponse)
async def jd_ingest(request: JDIngestRequest):
    """Ingest JD items into global FAISS index."""
    try:
        items = [item.model_dump() for item in request.items]
        count = rag.ingest_jds(items)
        return JDIngestResponse(status="ok", jd_count_added=count)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/resume/upload", response_model=ResumeUploadResponse)
async def resume_upload(
    file: UploadFile = File(None),
    text: str = Form(None),
    session_id: str = Form(None)
):
    """Upload resume via file or text. Returns upload_id for status tracking."""
    # Determine content
    resume_text = None
    
    if file:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")
        
        try:
            resume_text = rag.parse_file(content, file.filename or "unknown.txt")
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")
    elif text:
        resume_text = text
    else:
        raise HTTPException(status_code=400, detail="Either file or text must be provided")
    
    if not resume_text or not resume_text.strip():
        raise HTTPException(status_code=400, detail="No text content found in upload")
    
    # Generate session_id if not provided
    if not session_id:
        session_id = uuid.uuid4().hex[:8]
    
    try:
        # Start background processing
        upload_id = rag.start_resume_processing(session_id, resume_text)
        
        return ResumeUploadResponse(
            session_id=session_id,
            upload_id=upload_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/resume/upload/json", response_model=ResumeUploadResponse)
async def resume_upload_json(request: ResumeUploadTextRequest):
    """Upload resume via JSON body."""
    session_id = request.session_id or uuid.uuid4().hex[:8]
    
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="Text content is required")
    
    try:
        upload_id = rag.start_resume_processing(session_id, request.text)
        
        return ResumeUploadResponse(
            session_id=session_id,
            upload_id=upload_id
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/resume/status", response_model=ResumeStatusResponse)
async def resume_status(upload_id: str = Query(...)):
    """Get resume processing status."""
    status_info = rag.get_upload_status(upload_id)
    return ResumeStatusResponse(
        upload_id=upload_id,
        status=status_info["status"],
        detail=status_info.get("detail", "")
    )


@app.get("/resume/structured", response_model=ResumeStructuredResponse)
async def resume_structured(session_id: str = Query(...)):
    """Get structured resume blocks extracted from upload."""
    structured = rag.load_resume_structured(session_id)
    if not structured:
        raise HTTPException(status_code=404, detail=f"No structured resume for session {session_id}")
    return ResumeStructuredResponse(
        session_id=session_id,
        structured=ResumeInputStructured(**structured)
    )


@app.post("/resume/materials/add", response_model=ResumeUploadResponse)
async def resume_materials_add(
    session_id: str = Form(...),
    file: UploadFile = File(None),
    text: str = Form(None),
):
    """Add materials to existing session. Merges with resume, re-runs extract + cluster + index. Use /resume/status with returned upload_id."""
    new_text = None
    if file:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")
        try:
            new_text = rag.parse_file(content, file.filename or "unknown.txt")
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to parse file: {e}")
    elif text:
        new_text = text
    else:
        raise HTTPException(status_code=400, detail="Either file or text must be provided")
    if not new_text or not new_text.strip():
        raise HTTPException(status_code=400, detail="No text content in materials")
    try:
        upload_id = rag.start_add_materials(session_id, new_text.strip())
        return ResumeUploadResponse(session_id=session_id, upload_id=upload_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/resume/materials/add/json", response_model=ResumeUploadResponse)
async def resume_materials_add_json(request: AddMaterialsRequest):
    """Add materials as JSON. Merges with session resume, re-runs pipeline. Poll /resume/status with upload_id."""
    if not request.text or not request.text.strip():
        raise HTTPException(status_code=400, detail="Text is required")
    try:
        upload_id = rag.start_add_materials(request.session_id, request.text.strip())
        return ResumeUploadResponse(session_id=request.session_id, upload_id=upload_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/analyze/fit", response_model=AnalyzeFitResponse)
async def analyze_fit(request: AnalyzeFitRequest):
    """Analyze resume fit against JD requirements."""
    # Check session is ready
    if not rag.session_is_ready(request.session_id):
        raise HTTPException(status_code=400, detail=f"Session {request.session_id} resume not ready. Check /resume/status.")
    
    try:
        # Build query from target role
        query = f"Skills and experience for {request.target_role} role"
        
        # Retrieve resume evidence
        resume_chunks = rag.search_resume_index(request.session_id, query)
        
        # Retrieve JD evidence
        if request.use_curated_jd:
            jd_chunks = rag.search_jd_index(query, role=request.target_role)
            if not jd_chunks:
                raise HTTPException(status_code=400, detail="No curated JDs found. Ingest JDs first or provide jd_text.")
        elif request.jd_text:
            jd_chunks = rag.search_temp_jd(request.jd_text, query)
        else:
            raise HTTPException(status_code=400, detail="Either use_curated_jd=true or provide jd_text")
        
        # Build prompt and call Gemini
        user_prompt = build_fit_prompt(resume_chunks, jd_chunks, request.target_role)
        result = gemini_client.generate(ROLE_FIT_SYSTEM, user_prompt, ROLE_FIT_SCHEMA)
        
        content = result.get("content", {})
        if isinstance(content, str):
            content = {}
        
        # Build response
        recommended_roles = []
        for r in content.get("recommended_roles", []):
            recommended_roles.append(RoleRecommendation(
                role=r.get("role", request.target_role),
                score=min(1.0, max(0.0, r.get("score", 0.5))),
                reasons=r.get("reasons", [])
            ))
        
        if not recommended_roles:
            recommended_roles.append(RoleRecommendation(
                role=request.target_role,
                score=0.5,
                reasons=["Analysis based on provided evidence"]
            ))
        
        reqs = content.get("requirements", {})
        gap = content.get("gap", {})
        
        return AnalyzeFitResponse(
            recommended_roles=recommended_roles,
            requirements=Requirements(
                must_have=reqs.get("must_have", []),
                nice_to_have=reqs.get("nice_to_have", [])
            ),
            gap=GapAnalysis(
                matched=gap.get("matched", []),
                missing=gap.get("missing", []),
                ask_user_questions=gap.get("ask_user_questions", [])
            ),
            evidence=Evidence(
                resume_chunks=resume_chunks,
                jd_chunks=jd_chunks
            )
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/analyze/match-by-cluster", response_model=MatchByClusterResponse)
async def analyze_match_by_cluster(request: MatchByClusterRequest):
    """Per-cluster match % vs JD with evidence. Requires extraction + clusters (upload resume first)."""
    if not rag.session_is_ready(request.session_id):
        raise HTTPException(
            status_code=400,
            detail=f"Session {request.session_id} resume not ready. Check /resume/status.",
        )
    data = rag.load_clusters(request.session_id)
    if not data or not data.get("clusters"):
        raise HTTPException(
            status_code=400,
            detail="No clusters found. Upload resume and wait for processing, then retry.",
        )
    clusters = data["clusters"]
    meta = rag.get_all_resume_chunks(request.session_id)
    resume_chunks = [
        EvidenceChunk(chunk_id=m["chunk_id"], text=m["text"], source="resume", score=1.0)
        for m in meta
    ]
    query = "Job requirements, skills, and experience"
    if request.use_curated_jd:
        jd_chunks = rag.search_jd_index(query, role=None)
        if not jd_chunks:
            raise HTTPException(
                status_code=400,
                detail="No curated JDs found. Ingest JDs first or provide jd_text.",
            )
    elif request.jd_text:
        jd_chunks = rag.search_temp_jd(request.jd_text, query)
    else:
        raise HTTPException(
            status_code=400,
            detail="Either use_curated_jd=true or provide jd_text.",
        )
    user_prompt = build_match_by_cluster_prompt(clusters, resume_chunks, jd_chunks)
    debug_info = None
    if request.debug:
        debug_info = {
            "system_prompt": MATCH_BY_CLUSTER_SYSTEM,
            "user_prompt": user_prompt,
            "clusters": [
                {
                    "cluster_id": g.get("cluster_id"),
                    "cluster_label": g.get("cluster_label"),
                    "items_count": len(g.get("items", [])),
                    "items": [
                        {"id": i.get("id"), "label": i.get("label"), "text": (i.get("text") or "")[:300]}
                        for i in g.get("items", [])
                    ],
                }
                for g in clusters
            ],
            "resume_chunks": [
                {"chunk_id": c.chunk_id, "text": c.text[:500], "source": c.source, "score": c.score}
                for c in resume_chunks
            ],
            "jd_chunks": [
                {"chunk_id": c.chunk_id, "text": c.text[:500], "source": c.source, "score": c.score}
                for c in jd_chunks
            ],
        }
    result = gemini_client.generate(
        MATCH_BY_CLUSTER_SYSTEM, user_prompt, MATCH_BY_CLUSTER_SCHEMA
    )
    content = result.get("content") or {}
    if isinstance(content, str):
        content = {}
    r_map = {c.chunk_id: c for c in resume_chunks}
    j_map = {c.chunk_id: c for c in jd_chunks}
    cluster_matches = []
    for m in content.get("cluster_matches", []):
        cluster = m.get("cluster", "")
        match_pct = max(0.0, min(1.0, m.get("match_pct", 0.0)))
        r_ids = m.get("resume_chunk_ids", [])
        j_ids = m.get("jd_chunk_ids", [])
        r_ev = [r_map[x] for x in r_ids if x in r_map]
        j_ev = [j_map[x] for x in j_ids if x in j_map]
        cluster_matches.append(
            ClusterMatch(
                cluster=cluster,
                match_pct=match_pct,
                evidence=Evidence(resume_chunks=r_ev, jd_chunks=j_ev),
            )
        )
    overall = content.get("overall_match_pct")
    if overall is not None:
        overall = max(0.0, min(1.0, overall))
    return MatchByClusterResponse(
        cluster_matches=cluster_matches,
        overall_match_pct=overall,
        debug=debug_info,
    )


def generate_resume_internal(request: ResumeGenerateRequest) -> tuple:
    """Internal function to generate resume. Returns (structured, evidence)."""
    query = f"Complete background for {request.target_role} position"
    
    # Retrieve evidence
    resume_chunks = rag.search_resume_index(request.session_id, query)
    
    if request.use_curated_jd:
        jd_chunks = rag.search_jd_index(query, role=request.target_role)
    elif request.jd_text:
        jd_chunks = rag.search_temp_jd(request.jd_text, query)
    else:
        jd_chunks = []
    
    # Build prompt and generate
    user_prompt = build_generate_prompt(resume_chunks, jd_chunks, request.target_role)
    result = gemini_client.generate(RESUME_GENERATE_SYSTEM, user_prompt, RESUME_GENERATE_SCHEMA)
    
    content = result.get("content", {})
    if isinstance(content, str):
        content = {"education": [], "experience": [], "skills": [], "need_info": []}
    
    # Extract structured data
    education = content.get("education", [])
    experience = content.get("experience", [])
    skills = content.get("skills", [])
    need_info = content.get("need_info", [])
    
    structured = ResumeStructured(
        education=education,
        experience=experience,
        skills=skills
    )
    
    evidence = Evidence(
        resume_chunks=resume_chunks,
        jd_chunks=jd_chunks
    )
    
    return structured, need_info, evidence


def build_markdown_from_structured(structured: ResumeStructured) -> str:
    """Build markdown resume from structured data."""
    lines = []
    
    lines.append("## Education")
    if structured.education:
        for item in structured.education:
            lines.append(f"- {item}")
    else:
        lines.append("_No education information available._")
    
    lines.append("")
    lines.append("## Experience")
    if structured.experience:
        for item in structured.experience:
            lines.append(f"- {item}")
    else:
        lines.append("_No experience information available._")
    
    lines.append("")
    lines.append("## Skills")
    if structured.skills:
        lines.append(", ".join(structured.skills))
    else:
        lines.append("_No skills information available._")
    
    return "\n".join(lines)


@app.post("/resume/generate", response_model=ResumeGenerateResponse)
async def resume_generate(request: ResumeGenerateRequest):
    """Generate tailored resume based on evidence."""
    # Check session is ready
    if not rag.session_is_ready(request.session_id):
        raise HTTPException(status_code=400, detail=f"Session {request.session_id} resume not ready.")
    
    try:
        structured, need_info, evidence = generate_resume_internal(request)
        resume_markdown = build_markdown_from_structured(structured)
        
        return ResumeGenerateResponse(
            resume_markdown=resume_markdown,
            resume_structured=structured,
            need_info=need_info,
            evidence=evidence
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/resume/export_docx")
async def resume_export_docx(request: ResumeGenerateRequest):
    """Generate and export resume as .docx file."""
    from fastapi.responses import Response
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from io import BytesIO
    
    # Check session is ready
    if not rag.session_is_ready(request.session_id):
        raise HTTPException(status_code=400, detail=f"Session {request.session_id} resume not ready.")
    
    try:
        structured, need_info, _ = generate_resume_internal(request)
        
        # Create Word document
        doc = Document()
        
        # Title
        title = doc.add_heading("Resume", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Education section
        doc.add_heading("Education", level=1)
        if structured.education:
            for item in structured.education:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
        else:
            doc.add_paragraph("No education information available.", style='Normal')
        
        # Experience section
        doc.add_heading("Experience", level=1)
        if structured.experience:
            for item in structured.experience:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
        else:
            doc.add_paragraph("No experience information available.", style='Normal')
        
        # Skills section
        doc.add_heading("Skills", level=1)
        if structured.skills:
            skills_text = ", ".join(structured.skills)
            doc.add_paragraph(skills_text, style='Normal')
        else:
            doc.add_paragraph("No skills information available.", style='Normal')
        
        # Need Info section (if any)
        if need_info:
            doc.add_heading("Information Needed", level=1)
            for item in need_info:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(item)
                run.italic = True
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=resume.docx"}
        )
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =============================================================================
# CLUSTERING
# =============================================================================

@app.post("/experience/cluster", response_model=ClusterResponse)
async def cluster_experience(request: ClusterRequest):
    """
    Cluster user's experience items.

    Primary: When session_id is provided and session has extraction + clusters (from resume upload),
    returns role-based clusters (MLE/DS/SWE/QR/QD) with evidence.

    Fallback: When no extraction (stickers-only, pasted text, or old session), aggregates
    items + resume_text + optional session resume chunks into a single "All Experiences" cluster.
    """
    session_id = request.session_id or uuid.uuid4().hex[:8]

    # Primary: use stored clusters from resume extraction
    if request.session_id:
        data = rag.load_clusters(request.session_id)
        if data and data.get("clusters"):
            clusters = []
            for c in data["clusters"]:
                try:
                    g = ClusteredGroup.model_validate(c)
                    clusters.append(g)
                except Exception:
                    continue
            if clusters:
                total = data.get("total_items", sum(len(g.items) for g in clusters))
                return ClusterResponse(
                    session_id=session_id,
                    clusters=clusters,
                    total_items=total,
                    role_fit_distribution=data.get("role_fit_distribution"),
                )

    # Fallback: sticker/paste flow, single cluster
    all_items: list[ExperienceItem] = list(request.items)
    if request.resume_text and request.resume_text.strip():
        all_items.append(ExperienceItem(
            id=f"pasted_{uuid.uuid4().hex[:6]}",
            label="resume",
            text=request.resume_text.strip(),
            source="pasted_text",
        ))
    if request.session_id and rag.session_is_ready(request.session_id):
        for chunk in rag.get_all_resume_chunks(request.session_id):
            all_items.append(ExperienceItem(
                id=chunk.get("chunk_id", f"chunk_{uuid.uuid4().hex[:6]}"),
                label="resume",
                text=chunk.get("text", ""),
                source="uploaded_resume",
            ))
    single_cluster = ClusteredGroup(
        cluster_id="cluster_0",
        cluster_label="All Experiences",
        items=all_items,
        summary="All user experiences grouped together (no extraction; use resume upload for MLE/DS/SWE/QR/QD clusters)",
    )
    return ClusterResponse(
        session_id=session_id,
        clusters=[single_cluster],
        total_items=len(all_items),
        role_fit_distribution=None,
    )


if __name__ == "__main__":
    import os
    import uvicorn
    
    host = os.getenv("APP_HOST", "0.0.0.0")
    port = int(os.getenv("APP_PORT", "8000"))
    
    uvicorn.run(app, host=host, port=port)
